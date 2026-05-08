import os
import io
import json
import base64
import asyncio
import logging
import threading
from datetime import datetime, timezone
import anthropic
import requests
from http.server import HTTPServer, BaseHTTPRequestHandler
from telegram import Update
from telegram.ext import Application, MessageHandler, CommandHandler, filters, ContextTypes

TELEGRAM_TOKEN = os.environ["TELEGRAM_TOKEN"]
ANTHROPIC_API_KEY = os.environ["ANTHROPIC_API_KEY"]
COMPOSIO_API_KEY = os.environ["COMPOSIO_API_KEY"]
CLICKUP_API_KEY = os.environ["CLICKUP_API_KEY"]
ZOHO_CLIENT_ID = os.environ["ZOHO_CLIENT_ID"]
ZOHO_CLIENT_SECRET = os.environ["ZOHO_CLIENT_SECRET"]
ZOHO_REFRESH_TOKEN = os.environ.get("ZOHO_REFRESH_TOKEN", "")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

HAIKU = "claude-haiku-4-5"
SONNET = "claude-sonnet-4-6"

SYSTEM_PROMPT = """You are Ali's personal assistant. You have tools for ClickUp tasks and Zoho emails.
Use get_zoho_emails for latest emails, search_zoho_emails to find emails by keyword/sender/subject.
Use get_clickup_tasks for tasks, send_zoho_email to send, create_clickup_task to create tasks.
When asked about a specific person, company, or topic, always use search_zoho_emails.
When the user asks to analyze, review, compare, or summarize contracts/agreements/documents
for a specific company and topic (e.g. "analyze Dupilumab contract from Excellgene"), call
analyze_contracts with company and topic — do NOT use search_zoho_emails for that.
Be very concise. Plain text only. No markdown tables."""

CLICKUP_HEADERS = {"Authorization": CLICKUP_API_KEY}
zoho_refresh_token = ZOHO_REFRESH_TOKEN

CONTRACTS_MEMORY_PATH = os.environ.get("CONTRACTS_MEMORY_PATH", "contracts_memory.json")


def pick_model(text):
    heavy = ["draft", "write", "analyze", "analyse", "summarize", "summarise", "compose", "explain", "detail"]
    return SONNET if any(w in text.lower() for w in heavy) else HAIKU


class TokenHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        global zoho_refresh_token
        if self.path.startswith("/zoho-auth"):
            from urllib.parse import urlparse, parse_qs
            params = parse_qs(urlparse(self.path).query)
            code = params.get("code", [None])[0]
            if code:
                resp = requests.post("https://accounts.zoho.com/oauth/v2/token", data={
                    "code": code,
                    "client_id": ZOHO_CLIENT_ID,
                    "client_secret": ZOHO_CLIENT_SECRET,
                    "redirect_uri": f"https://{self.headers.get('Host')}/zoho-auth",
                    "grant_type": "authorization_code"
                }).json()
                rt = resp.get("refresh_token", "")
                if rt:
                    zoho_refresh_token = rt
                    msg = f"SUCCESS! ZOHO_REFRESH_TOKEN:\n\n{rt}"
                else:
                    msg = f"Error: {resp}"
            else:
                msg = "No code found."
            self.send_response(200)
            self.send_header("Content-type", "text/plain")
            self.end_headers()
            self.wfile.write(msg.encode())
        else:
            self.send_response(200)
            self.send_header("Content-type", "text/plain")
            self.end_headers()
            self.wfile.write(b"Bot is running!")
    def log_message(self, format, *args):
        pass


def start_web_server():
    port = int(os.environ.get("PORT", 8080))
    HTTPServer(("0.0.0.0", port), TokenHandler).serve_forever()


def get_zoho_token():
    if not zoho_refresh_token:
        return None
    resp = requests.post("https://accounts.zoho.com/oauth/v2/token", data={
        "refresh_token": zoho_refresh_token,
        "client_id": ZOHO_CLIENT_ID,
        "client_secret": ZOHO_CLIENT_SECRET,
        "grant_type": "refresh_token"
    }).json()
    return resp.get("access_token")


def get_zoho_account():
    token = get_zoho_token()
    if not token:
        return None, None, None
    accounts = requests.get(
        "https://mail.zoho.com/api/accounts",
        headers={"Authorization": f"Zoho-oauthtoken {token}"}
    ).json().get("data", [])
    if not accounts:
        return None, None, None
    aid = accounts[0].get("accountId")
    from_addr = accounts[0].get("primaryEmailAddress") or accounts[0].get("emailAddress")
    return token, aid, from_addr


def format_messages(msgs):
    if not msgs:
        return "No emails found."
    return "\n".join([
        f"{i+1}. From: {m.get('fromAddress','?')} | {m.get('subject','(no subject)')}"
        for i, m in enumerate(msgs)
    ])


def get_zoho_emails():
    try:
        token, aid, _ = get_zoho_account()
        if not token:
            return "Zoho not connected."
        msgs = requests.get(
            f"https://mail.zoho.com/api/accounts/{aid}/messages/view?limit=5&sortorder=false",
            headers={"Authorization": f"Zoho-oauthtoken {token}"}
        ).json().get("data", [])
        return format_messages(msgs)
    except Exception as e:
        return f"Error: {e}"


def search_zoho_emails(query):
    try:
        token, aid, _ = get_zoho_account()
        if not token:
            return "Zoho not connected."
        q = query.strip()
        if "@" in q and " " not in q:
            search_key = f"sender:{q}"
        else:
            value = f'"{q}"' if " " in q else q
            search_key = f"entire:{value}"
        msgs = requests.get(
            f"https://mail.zoho.com/api/accounts/{aid}/messages/search",
            headers={"Authorization": f"Zoho-oauthtoken {token}"},
            params={"searchKey": search_key, "limit": 10, "sortorder": "false"}
        ).json().get("data", [])
        return format_messages(msgs)
    except Exception as e:
        return f"Error: {e}"


def send_zoho_email(to, subject, body):
    try:
        token, aid, from_addr = get_zoho_account()
        if not token:
            return "Zoho not connected."
        resp = requests.post(
            f"https://mail.zoho.com/api/accounts/{aid}/messages",
            headers={"Authorization": f"Zoho-oauthtoken {token}"},
            json={"fromAddress": from_addr, "toAddress": to, "subject": subject, "content": body, "mailFormat": "plaintext"}
        ).json()
        if resp.get("status", {}).get("code") == 200:
            return f"Email sent to {to}."
        return f"Send error: {resp}"
    except Exception as e:
        return f"Error: {e}"


def get_clickup_tasks(search=""):
    try:
        teams = requests.get("https://api.clickup.com/api/v2/team", headers=CLICKUP_HEADERS).json()
        tid = teams["teams"][0]["id"]
        url = f"https://api.clickup.com/api/v2/team/{tid}/task?include_closed=false&page=0"
        if search:
            url += f"&search={search}"
        tasks = requests.get(url, headers=CLICKUP_HEADERS).json().get("tasks", [])[:5]
        if not tasks:
            return "No tasks found."
        return "\n".join([
            f"- {t.get('name','?')} [{t.get('status',{}).get('status','?')}] [{t.get('priority',{}).get('priority','-') if t.get('priority') else '-'}]"
            for t in tasks
        ])
    except Exception as e:
        return f"Error: {e}"


def create_clickup_task(name, list_id, description=""):
    try:
        r = requests.post(
            f"https://api.clickup.com/api/v2/list/{list_id}/task",
            headers=CLICKUP_HEADERS,
            json={"name": name, "description": description}
        ).json()
        return f"Created: {r.get('name', 'done')}"
    except Exception as e:
        return f"Error: {e}"


def _slug(s):
    return "".join(c.lower() if c.isalnum() else "_" for c in (s or "")).strip("_") or "unknown"


def load_contracts_memory():
    if not os.path.exists(CONTRACTS_MEMORY_PATH):
        return {"contracts": {}}
    try:
        with open(CONTRACTS_MEMORY_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        if "contracts" not in data:
            data["contracts"] = {}
        return data
    except Exception as e:
        logger.warning(f"contracts_memory unreadable, starting fresh: {e}")
        return {"contracts": {}}


def save_contracts_memory(data):
    tmp = CONTRACTS_MEMORY_PATH + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    os.replace(tmp, CONTRACTS_MEMORY_PATH)


def list_zoho_attachments(token, aid, folder_id, msg_id):
    r = requests.get(
        f"https://mail.zoho.com/api/accounts/{aid}/folders/{folder_id}/messages/{msg_id}/attachmentinfo",
        headers={"Authorization": f"Zoho-oauthtoken {token}"}
    ).json()
    return (r.get("data") or {}).get("attachments") or []


def download_zoho_attachment(token, aid, folder_id, msg_id, att_id):
    r = requests.get(
        f"https://mail.zoho.com/api/accounts/{aid}/folders/{folder_id}/messages/{msg_id}/attachments/{att_id}",
        headers={"Authorization": f"Zoho-oauthtoken {token}"}
    )
    r.raise_for_status()
    return r.content


def extract_docx_text(data):
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            line = "\t".join((c.text or "").strip() for c in row.cells)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def extract_xlsx_text(data):
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            line = "\t".join("" if v is None else str(v) for v in row)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def build_attachment_block(filename, data):
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": base64.standard_b64encode(data).decode("ascii"),
            },
        }
    if name.endswith(".docx"):
        text = extract_docx_text(data)
        return {"type": "text", "text": f"[Document: {filename}]\n\n{text}"}
    if name.endswith((".xlsx", ".xlsm")):
        text = extract_xlsx_text(data)
        return {"type": "text", "text": f"[Spreadsheet: {filename}]\n\n{text}"}
    return None


ANALYSIS_INSTRUCTIONS = (
    "You are reviewing a document attached to an email thread between Ali and {company} "
    "about the topic '{topic}'. Extract a structured summary of this document.\n\n"
    "Document filename: {filename}\n"
    "Version number to assign: {version}\n\n"
    "Previous version summary (use to compute diff; null means this is the first version):\n"
    "{prev_summary}\n\n"
    "Return STRICT JSON only — no prose, no markdown fences. Schema:\n"
    "{{\n"
    '  "date": "<ISO date the document is dated, or today if none>",\n'
    '  "key_terms": [<short bullet strings, max 10>],\n'
    '  "open_issues": [<strings>],\n'
    '  "resolved_issues": [<strings — issues from previous version that this version resolves; [] if no previous>],\n'
    '  "changes_from_previous": [<strings — concrete diffs from previous version>] | null\n'
    "}}\n\n"
    "If there is no previous version, set changes_from_previous to null and resolved_issues to []."
)


def analyze_attachment_with_claude(company, topic, filename, version, prev_summary, content_block):
    prev_compact = None
    if prev_summary:
        prev_compact = {
            "version": prev_summary.get("version"),
            "date": prev_summary.get("date"),
            "key_terms": prev_summary.get("key_terms"),
            "open_issues": prev_summary.get("open_issues"),
        }
    instructions = ANALYSIS_INSTRUCTIONS.format(
        company=company,
        topic=topic,
        filename=filename,
        version=version,
        prev_summary=json.dumps(prev_compact, ensure_ascii=False) if prev_compact else "null",
    )
    resp = claude_client.messages.create(
        model=SONNET,
        max_tokens=2000,
        messages=[{"role": "user", "content": [content_block, {"type": "text", "text": instructions}]}],
    )
    text = resp.content[0].text.strip()
    if text.startswith("```"):
        text = text.split("```", 2)[1]
        if text.lower().startswith("json"):
            text = text[4:]
        text = text.strip().rsplit("```", 1)[0].strip()
    return json.loads(text)


def analyze_contracts(company, topic):
    token, aid, _ = get_zoho_account()
    if not token:
        return "Zoho not connected.", None

    search_key = f'entire:"{company}"::entire:"{topic}"'
    msgs = requests.get(
        f"https://mail.zoho.com/api/accounts/{aid}/messages/search",
        headers={"Authorization": f"Zoho-oauthtoken {token}"},
        params={"searchKey": search_key, "limit": 50, "sortorder": "true"},
    ).json().get("data") or []

    if not msgs:
        return f"No emails found for {company} / {topic}.", None

    memory = load_contracts_memory()
    key = f"{_slug(company)}__{_slug(topic)}"
    record = memory["contracts"].get(key) or {
        "company": company,
        "topic": topic,
        "current_version": 0,
        "last_updated": None,
        "processed_attachments": [],
        "versions": [],
    }
    processed = set(record["processed_attachments"])

    new_count = 0
    skipped = []
    failed = []

    for m in msgs:
        if str(m.get("hasAttachment", "0")) != "1":
            continue
        msg_id = m.get("messageId")
        folder_id = m.get("folderId")
        if not msg_id or not folder_id:
            continue
        try:
            atts = list_zoho_attachments(token, aid, folder_id, msg_id)
        except Exception as e:
            logger.warning(f"attachmentinfo failed for {msg_id}: {e}")
            continue
        for a in atts:
            att_id = a.get("attachmentId")
            name = a.get("attachmentName", "")
            if not att_id:
                continue
            key_id = f"{msg_id}:{att_id}"
            if key_id in processed:
                continue
            ext_supported = name.lower().endswith((".pdf", ".docx", ".xlsx", ".xlsm"))
            if not ext_supported:
                skipped.append(name)
                processed.add(key_id)
                continue
            try:
                data = download_zoho_attachment(token, aid, folder_id, msg_id, att_id)
                block = build_attachment_block(name, data)
                if block is None:
                    skipped.append(name)
                    processed.add(key_id)
                    continue
                prev = record["versions"][-1] if record["versions"] else None
                next_version = record["current_version"] + 1
                result = analyze_attachment_with_claude(
                    company, topic, name, next_version, prev, block
                )
            except Exception as e:
                logger.exception(f"analysis failed for {name}")
                failed.append(f"{name}: {str(e)[:80]}")
                continue
            result["version"] = next_version
            result["company"] = company
            result["topic"] = topic
            result["source"] = {"message_id": msg_id, "attachment_id": att_id, "filename": name}
            record["versions"].append(result)
            record["current_version"] = next_version
            record["last_updated"] = datetime.now(timezone.utc).strftime("%Y-%m-%d")
            processed.add(key_id)
            new_count += 1

    record["processed_attachments"] = sorted(processed)
    memory["contracts"][key] = record
    save_contracts_memory(memory)

    return None, {
        "record": record,
        "new_count": new_count,
        "skipped": skipped,
        "failed": failed,
    }


def format_contract_comparison(record, new_count, skipped, failed):
    versions = record.get("versions") or []
    company = record.get("company", "?")
    topic = record.get("topic", "?")
    if not versions:
        return f"{company} — {topic}\nNo supported attachments analyzed."

    latest = versions[-1]
    prev = versions[-2] if len(versions) > 1 else None

    lines = [
        f"{company} — {topic}",
        f"Versions: {len(versions)} | New this run: {new_count} | Last updated: {record.get('last_updated','?')}",
        "",
        f"--- Latest (v{latest.get('version','?')}) ---",
        f"Date: {latest.get('date','?')}",
        f"Source: {(latest.get('source') or {}).get('filename','?')}",
    ]
    key_terms = latest.get("key_terms") or []
    if key_terms:
        lines.append("Key terms:")
        for t in key_terms[:10]:
            lines.append(f"  • {t}")
    open_issues = latest.get("open_issues") or []
    if open_issues:
        lines.append("Open issues:")
        for i in open_issues[:10]:
            lines.append(f"  • {i}")

    if prev:
        lines += [
            "",
            f"--- Changes v{prev.get('version','?')} → v{latest.get('version','?')} ---",
        ]
        for c in (latest.get("changes_from_previous") or [])[:15]:
            lines.append(f"  • {c}")
        resolved = latest.get("resolved_issues") or []
        if resolved:
            lines.append("Resolved:")
            for r in resolved[:10]:
                lines.append(f"  • {r}")

    if skipped:
        lines += ["", f"Skipped (unsupported): {', '.join(skipped[:5])}"]
    if failed:
        lines += ["", f"Failed: {', '.join(failed[:3])}"]

    text = "\n".join(lines)
    if len(text) > 3990:
        text = text[:3980] + "\n…"
    return text


tools = [
    {
        "name": "get_zoho_emails",
        "description": "Get Ali's 5 most recent Zoho emails",
        "input_schema": {"type": "object", "properties": {}}
    },
    {
        "name": "search_zoho_emails",
        "description": "Search Zoho emails by sender name, company, or subject keyword",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search term e.g. Excellgene, Benjamin, contract"}
            },
            "required": ["query"]
        }
    },
    {
        "name": "send_zoho_email",
        "description": "Send an email via Zoho",
        "input_schema": {
            "type": "object",
            "properties": {
                "to": {"type": "string"},
                "subject": {"type": "string"},
                "body": {"type": "string"}
            },
            "required": ["to", "subject", "body"]
        }
    },
    {
        "name": "get_clickup_tasks",
        "description": "Get Ali's ClickUp tasks, filtered by optional keyword",
        "input_schema": {
            "type": "object",
            "properties": {
                "search": {"type": "string"}
            }
        }
    },
    {
        "name": "create_clickup_task",
        "description": "Create a new ClickUp task",
        "input_schema": {
            "type": "object",
            "properties": {
                "name": {"type": "string"},
                "description": {"type": "string"},
                "list_id": {"type": "string"}
            },
            "required": ["name", "list_id"]
        }
    },
    {
        "name": "analyze_contracts",
        "description": (
            "Analyze contract/agreement documents (PDF, DOCX, XLSX) attached to "
            "Zoho emails between Ali and a specific company on a specific topic. "
            "Returns a versioned summary with key terms, open issues, and a diff "
            "from the previous version. Use this whenever the user asks to "
            "analyze/review/compare/summarize contracts or documents for a "
            "company+topic pair."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "company": {"type": "string", "description": "Company name, e.g. Excellgene"},
                "topic": {"type": "string", "description": "Contract topic, e.g. Dupilumab"}
            },
            "required": ["company", "topic"]
        }
    }
]

conversation_histories = {}


def _block_type(block):
    if isinstance(block, dict):
        return block.get("type")
    return getattr(block, "type", None)


def clean_orphan_tool_uses(messages):
    cleaned = []
    i = 0
    while i < len(messages):
        msg = messages[i]
        content = msg.get("content")
        has_tool_use = (
            msg.get("role") == "assistant"
            and isinstance(content, list)
            and any(_block_type(b) == "tool_use" for b in content)
        )
        if has_tool_use:
            nxt = messages[i + 1] if i + 1 < len(messages) else None
            nxt_has_tool_result = (
                nxt is not None
                and nxt.get("role") == "user"
                and isinstance(nxt.get("content"), list)
                and any(_block_type(b) == "tool_result" for b in nxt["content"])
            )
            if not nxt_has_tool_result:
                i += 1
                continue
        cleaned.append(msg)
        i += 1
    return cleaned


def get_claude_response(messages, model):
    response = claude_client.messages.create(
        model=model,
        max_tokens=400,
        system=SYSTEM_PROMPT,
        messages=clean_orphan_tool_uses(messages),
        tools=tools
    )
    if response.stop_reason == "tool_use":
        tool_use = next(b for b in response.content if b.type == "tool_use")
        if tool_use.name == "get_zoho_emails":
            result = get_zoho_emails()
        elif tool_use.name == "search_zoho_emails":
            result = search_zoho_emails(**tool_use.input)
        elif tool_use.name == "send_zoho_email":
            result = send_zoho_email(**tool_use.input)
        elif tool_use.name == "get_clickup_tasks":
            result = get_clickup_tasks(**tool_use.input)
        elif tool_use.name == "create_clickup_task":
            result = create_clickup_task(**tool_use.input)
        elif tool_use.name == "analyze_contracts":
            err, data = analyze_contracts(**tool_use.input)
            if err:
                return err
            return format_contract_comparison(
                data["record"], data["new_count"], data["skipped"], data["failed"]
            )
        else:
            result = "Unknown tool"
        messages = messages + [
            {"role": "assistant", "content": response.content},
            {"role": "user", "content": [{"type": "tool_result", "tool_use_id": tool_use.id, "content": result}]}
        ]
        response = claude_client.messages.create(
            model=model,
            max_tokens=400,
            system=SYSTEM_PROMPT,
            messages=clean_orphan_tool_uses(messages),
            tools=tools
        )
    return next((b.text for b in response.content if hasattr(b, 'text')), 'Done.')


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("👋 Hi Ali! Ask me about your emails or ClickUp tasks.")


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_message = update.message.text
    user_id = str(update.effective_user.id)
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")

    if user_id not in conversation_histories:
        conversation_histories[user_id] = []

    conversation_histories[user_id].append({"role": "user", "content": user_message})

    if len(conversation_histories[user_id]) > 6:
        conversation_histories[user_id] = conversation_histories[user_id][-6:]

    model = pick_model(user_message)
    logger.info(f"Model: {model} | Message: {user_message[:50]}")

    try:
        reply = get_claude_response(conversation_histories[user_id], model)
        conversation_histories[user_id].append({"role": "assistant", "content": reply})
        await update.message.reply_text(reply)
    except Exception as e:
        logger.error(f"Error: {e}")
        await update.message.reply_text(f"Error: {str(e)[:200]}")


async def clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    conversation_histories.pop(user_id, None)
    conversation_histories[user_id] = []
    await update.message.reply_text("✅ Cleared.")


async def analyze_contracts_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    args = context.args or []
    if len(args) < 2:
        await update.message.reply_text(
            "Usage: /analyze_contracts <company> <topic>\n"
            "Examples:\n"
            "  /analyze_contracts Excellgene Dupilumab\n"
            "  /analyze_contracts Sartorius equipment\n"
            "  /analyze_contracts WHO regulatory"
        )
        return
    company = args[0]
    topic = " ".join(args[1:])
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
    await update.message.reply_text(f"Analyzing {company} / {topic}… this may take a minute.")
    try:
        err, result = await asyncio.to_thread(analyze_contracts, company, topic)
        if err:
            await update.message.reply_text(err)
            return
        text = format_contract_comparison(
            result["record"], result["new_count"], result["skipped"], result["failed"]
        )
        await update.message.reply_text(text)
    except Exception as e:
        logger.exception("analyze_contracts_cmd failed")
        await update.message.reply_text(f"Error: {str(e)[:200]}")


def main():
    threading.Thread(target=start_web_server, daemon=True).start()
    app = Application.builder().token(TELEGRAM_TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("clear", clear))
    app.add_handler(CommandHandler("analyze_contracts", analyze_contracts_cmd))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    logger.info("Bot running...")
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
